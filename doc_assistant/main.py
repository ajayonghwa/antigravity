import os
import json
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches
from ollama_client import OllamaClient
from pdf_utils import load_all_references

class DocAssistant:
    def __init__(self, config_path="config.json"):
        self.load_config(config_path)
        self.client = OllamaClient(model=self.model)
        self.root = Path(".").resolve()
        self.paths = {
            "refs": self.root / "references",
            "images": self.root / "images",
            "content": self.root / "content",
            "output": self.root / "output"
        }
        for p in self.paths.values():
            p.mkdir(exist_ok=True)

    def load_config(self, path):
        if os.path.exists(path):
            with open(path, "r") as f:
                config = json.load(f)
        else:
            config = {"model": "gemma2:9b", "password": None} # Default to gemma2:9b as gemma4 is likely gemma2
        self.model = config.get("model", "gemma2:9b")
        self.password = config.get("password", None)

    def step1_generate_outline(self, topic):
        print(f"\n[Step 1] '{topic}' 주제로 개요를 생성합니다...")
        refs = load_all_references(self.paths["refs"], self.password)
        
        prompt = f"주제: {topic}\n\n참고 문헌 내용:\n{refs[:4000]}\n\n위 내용을 바탕으로 문서의 개요를 짜주세요. 각 섹션은 'Section X: 제목' 형식으로 작성하고, 간단한 설명도 포함해주세요."
        system_prompt = "당신은 전문적인 문서 작성을 돕는 어시스턴트입니다. 사용자가 제공한 주제와 참고 자료를 바탕으로 논리적이고 체계적인 목차(개요)를 생성하세요. 한국어로 답변하세요."
        
        outline = self.client.generate(prompt, system_prompt)
        
        outline_path = self.paths["content"] / "outline.md"
        outline_path.write_text(outline, encoding="utf-8")
        print(f"개요가 생성되었습니다: {outline_path}")
        return outline

    def step2_interactive_draft(self):
        outline_path = self.paths["content"] / "outline.md"
        drafts_path = self.paths["content"] / "drafts.json"
        
        if not outline_path.exists():
            print("개요 파일이 없습니다. Step 1을 먼저 실행하세요.")
            return

        # 1. 안내 문구 출력
        print("\n" + "="*50)
        print(" [Step 2] 대화형 초안 작성 가이드")
        print(" - 섹션별 질문에 답변을 입력하면 AI가 전문적인 문장으로 다듬어줍니다.")
        print(" - 'q' 또는 'exit' 입력: 작업 중단 및 현재까지 내용 저장")
        print(" - 그냥 Enter: 해당 섹션 내용 없이 건너뛰기")
        print("="*50)

        # 2. 기존 진행 상황 확인
        drafts = {}
        if drafts_path.exists():
            with open(drafts_path, "r", encoding="utf-8") as f:
                drafts = json.load(f)
            
            completed_secs = [s for s, c in drafts.items() if c]
            print(f"\n현재 진행 상황: 총 {len(completed_secs)}개 섹션 작성 완료")
            for i, sec in enumerate(completed_secs):
                print(f"  [{i+1}] {sec}")
            
            print("\n작업을 어떻게 진행할까요?")
            print("1. 이어서 하기 (작성된 부분 건너뛰기)")
            print("2. 처음부터 다시 하기 (기존 내용 초기화)")
            print("3. 특정 섹션부터 다시 하기")
            
            mode = input("선택 (1/2/3): ")
            if mode == "2":
                drafts = {}
                resume_from = None
            elif mode == "3":
                idx = int(input("다시 시작할 섹션 번호를 선택하세요: ")) - 1
                # 선택한 번호 이후의 데이터 삭제
                keys = list(drafts.keys())
                for k in keys[idx:]:
                    del drafts[k]
                resume_from = None
            else:
                resume_from = "auto"
        else:
            resume_from = None

        outline_text = outline_path.read_text(encoding="utf-8")
        sections = [line.strip() for line in outline_text.split("\n") if "Section" in line]
        
        for section in sections:
            # 이어하기 로직
            if resume_from == "auto" and section in drafts and drafts[section]:
                print(f"\n[Skip] {section} (이미 작성됨)")
                continue

            print(f"\n>>> {section}")
            user_input = input("내용 입력: ")
            
            if user_input.lower() in ["q", "exit", "quit"]:
                print("\n작업을 중단합니다. 현재까지의 내용을 저장했습니다.")
                break

            if user_input.strip():
                print("...AI가 문장을 다듬는 중입니다...")
                refining_prompt = f"섹션: {section}\n사용자 입력: {user_input}\n\n이 내용을 바탕으로 공식적인 문서에 적합한 문단으로 다듬어주세요."
                refined_text = self.client.generate(refining_prompt, "당신은 구어체 답변을 전문적인 문서 스타일로 변환하는 에디터입니다.")
                
                if "Error" in refined_text:
                    print(f"⚠️ 오류 발생: {refined_text}")
                    print("이 섹션을 다시 시도하거나 나중에 수정해 주세요.")
                else:
                    drafts[section] = refined_text
                    print(f"--- 다듬어진 내용 ---\n{refined_text}")
            else:
                if section not in drafts:
                    drafts[section] = ""

            # 매 섹션마다 자동 저장
            with open(drafts_path, "w", encoding="utf-8") as f:
                json.dump(drafts, f, ensure_ascii=False, indent=2)

        print(f"\n현재까지 {len(drafts)}개 섹션의 초안이 저장되었습니다.")

    def step4_manual_edit(self):
        print("\n[Step 4] 개요 및 내용 수동 수정")
        print("1. 새로운 섹션 추가")
        print("2. 특정 섹션 내용 직접 수정")
        print("3. 메인 메뉴로 돌아가기")
        
        choice = input("선택하세요: ")
        drafts_path = self.paths["content"] / "drafts.json"
        
        drafts = {}
        if drafts_path.exists():
            with open(drafts_path, "r", encoding="utf-8") as f:
                drafts = json.load(f)

        if choice == "1":
            new_sec = input("추가할 섹션 제목 (예: Section 5: 결론): ")
            new_content = input("섹션 내용 (직접 입력): ")
            drafts[new_sec] = new_content
            print("섹션이 추가되었습니다.")
        elif choice == "2":
            if not drafts:
                print("수정할 내용이 없습니다.")
                return
            for i, sec in enumerate(drafts.keys()):
                print(f"{i+1}. {sec}")
            idx = int(input("수정할 번호 선택: ")) - 1
            target_sec = list(drafts.keys())[idx]
            print(f"\n[현재 내용]: {drafts[target_sec]}")
            new_content = input("새로운 내용 (변경 없으면 Enter): ")
            if new_content.strip():
                drafts[target_sec] = new_content
                print("수정 완료.")
        else:
            return

        with open(drafts_path, "w", encoding="utf-8") as f:
            json.dump(drafts, f, ensure_ascii=False, indent=2)

    def step3_assemble(self, filename="Final_Document"):
        print("\n[Step 3] 이미지를 포함하여 최종 문서를 생성합니다...")
        drafts_path = self.paths["content"] / "drafts.json"
        if not drafts_path.exists():
            print("초안 데이터가 없습니다. Step 2를 완료하세요.")
            return

        with open(drafts_path, "r", encoding="utf-8") as f:
            drafts = json.load(f)

        # 1. Markdown 파일 먼저 생성
        md_content = "# 문서 제목\n\n"
        for section, content in drafts.items():
            md_content += f"## {section}\n\n{content}\n\n"
            
            # 이미지 매칭 (섹션 번호 또는 이름 키워드)
            section_num = re.search(r'\d+', section)
            
            for img_folder in self.paths["images"].iterdir():
                if not img_folder.is_dir(): continue
                
                folder_num = re.search(r'\d+', img_folder.name)
                # 번호가 일치하거나 폴더명이 섹션 제목에 포함된 경우
                is_match = (section_num and folder_num and section_num.group() == folder_num.group()) or (img_folder.name in section)
                
                if is_match:
                    for img_file in img_folder.glob("*"):
                        if img_file.suffix.lower() in [".png", ".jpg", ".jpeg"]:
                            md_content += f"![{img_file.name}](../images/{img_folder.name}/{img_file.name})\n\n"
        
        md_save_path = self.paths["output"] / f"{filename}.md"
        md_save_path.write_text(md_content, encoding="utf-8")
        print(f"Markdown 파일이 생성되었습니다: {md_save_path}")

        # 2. DOCX 파일 생성
        doc = Document()
        doc.add_heading("문서 제목", 0)

        for section, content in drafts.items():
            doc.add_heading(section, level=1)
            doc.add_paragraph(content)

            section_num = re.search(r'\d+', section)
            for img_folder in self.paths["images"].iterdir():
                if not img_folder.is_dir(): continue
                folder_num = re.search(r'\d+', img_folder.name)
                is_match = (section_num and folder_num and section_num.group() == folder_num.group()) or (img_folder.name in section)
                
                if is_match:
                    doc.add_paragraph(f"--- 관련 이미지: {img_folder.name} ---")
                    for img_file in sorted(img_folder.glob("*")):
                        if img_file.suffix.lower() in [".png", ".jpg", ".jpeg"]:
                            try:
                                doc.add_picture(str(img_file), width=Inches(5))
                                doc.add_paragraph(f"그림: {img_file.name}")
                            except Exception as e:
                                print(f"이미지 삽입 실패 ({img_file.name}): {e}")

        docx_save_path = self.paths["output"] / f"{filename}.docx"
        doc.save(docx_save_path)
        print(f"Word 문서가 생성되었습니다: {docx_save_path}")

if __name__ == "__main__":
    assistant = DocAssistant()
    
    print("=== 오프라인 문서 작성 도우미 ===")
    print("1. 개요 생성 (Step 1)")
    print("2. 대화형 내용 작성 (Step 2) [이어하기/중단 지원]")
    print("3. 최종 문서 생성 (.md & .docx) (Step 3)")
    print("4. 개요 및 내용 수동 수정 (Step 4)")
    
    choice = input("명령을 선택하세요 (1/2/3/4): ")
    
    if choice == "1":
        topic = input("작성할 문서의 주제를 입력하세요: ")
        assistant.step1_generate_outline(topic)
    elif choice == "2":
        assistant.step2_interactive_draft()
    elif choice == "3":
        assistant.step3_assemble()
    elif choice == "4":
        assistant.step4_manual_edit()
    else:
        print("잘못된 선택입니다.")
